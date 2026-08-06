"""
Parser de DOCX com extração de metadados (core.xml)
"""
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, Any
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from engine.logger import ATSLogger


class DOCXParser:
    def __init__(self, logger: ATSLogger = None):
        self.logger = logger

    def parse(self, file_path: str) -> Dict[str, Any]:
        if self.logger:
            self.logger.log("PARSE", "Iniciando parsing de DOCX", {"file": file_path})

        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx não instalado. Execute: pip install python-docx")

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_text = "\n".join(paragraphs)

        # Extrai tabelas também
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                tables_text.append(row_text)

        if tables_text:
            raw_text += "\n\n[TABELAS]\n" + "\n".join(tables_text)

        raw_text = self._clean_text(raw_text)

        # === EXTRAÇÃO DE METADADOS (core.xml) ===
        metadata = self._extract_metadata(file_path)

        if self.logger:
            self.logger.log("PARSE", "DOCX parseado com sucesso", {
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "chars": len(raw_text),
                "metadata": metadata
            })

        return {
            "raw_text": raw_text,
            "pages": [raw_text],
            "metadata": {
                **metadata,
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "is_image_pdf": False,
            },
            "file_type": "docx"
        }

    def _extract_metadata(self, file_path: str) -> Dict[str, str]:
        """Extrai metadados do docProps/core.xml do .docx (ZIP)"""
        metadata = {
            "title": "",
            "author": "",
            "subject": "",
            "keywords": "",
            "description": "",
        }
        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                if 'docProps/core.xml' in z.namelist():
                    core_xml = z.read('docProps/core.xml').decode('utf-8')
                    # Parse XML com namespace
                    ns = {
                        'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
                        'dc': 'http://purl.org/dc/elements/1.1/',
                        'dcterms': 'http://purl.org/dc/terms/',
                    }
                    root = ET.fromstring(core_xml)

                    def get(tag, ns_key='dc'):
                        elem = root.find(f'{ns_key}:{tag}', ns)
                        return elem.text.strip() if elem is not None and elem.text else ""

                    metadata["title"] = get("title")
                    metadata["author"] = get("creator", "cp") or get("author")
                    metadata["subject"] = get("subject")
                    metadata["keywords"] = get("keywords", "cp")
                    metadata["description"] = get("description")
        except Exception as e:
            if self.logger:
                self.logger.log("PARSE", "Erro ao extrair metadados DOCX", {"error": str(e)}, severity="WARN")
        return metadata

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        return text.strip()
